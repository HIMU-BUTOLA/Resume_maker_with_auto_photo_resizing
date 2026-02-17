import cv2
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_bottom_border(paragraph):
    """
    Adds a bottom border to a paragraph (used for headings).
    Requires manipulating the Oxml (XML) of the document.
    """
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    
    # Border attributes
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')      # Line width (1/8 pt)
    bottom.set(qn('w:space'), '1')   # Space between text and border
    bottom.set(qn('w:color'), 'auto') # Color
    
    pBdr.append(bottom)
    pPr.append(pBdr)

def process_passport_photo(input_path, output_path):
    """
    Detects a face, adds padding, crops to 3.5:4.5 ratio, saves file.
    """
    if not os.path.exists(input_path):
        print(f"Error: Image {input_path} not found.")
        return False

    img = cv2.imread(input_path)
    if img is None:
        return False

    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
    faces = face_cascade.detectMultiScale(gray, 1.3, 5)

    if len(faces) > 0:
        x, y, w, h = faces[0]
        pad_top = int(h * 0.5)
        pad_bottom = int(h * 0.5)
        pad_left = int(w * 0.3)
        pad_right = int(w * 0.3)

        y1 = max(0, y - pad_top)
        y2 = min(img.shape[0], y + h + pad_bottom)
        x1 = max(0, x - pad_left)
        x2 = min(img.shape[1], x + w + pad_right)
        
        crop_img = img[y1:y2, x1:x2]
    else:
        h, w = img.shape[:2]
        crop_img = img[0:h, int(w*0.25):int(w*0.75)]

    # Resize to standard passport ratio
    final_img = cv2.resize(crop_img, (413, 531))
    cv2.imwrite(output_path, final_img)
    return True

def generate_resume(personal_info, photo_path):
    doc = Document()

    # --- 1. Header Section ---
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(4.5)
    table.columns[1].width = Inches(1.5)

    # Name and Contact
    cell_info = table.cell(0, 0)
    name_para = cell_info.paragraphs[0]
    run = name_para.add_run(personal_info.get("Full Name", ""))
    run.bold = True
    run.font.size = Pt(24)
    
    contact_keys = ["Email", "Phone", "Address", "LinkedIn"]
    for key in contact_keys:
        val = personal_info.get(key, "NA")
        if val != "NA":
            cell_info.add_paragraph(f"{val}")

    # Photo
    cell_photo = table.cell(0, 1)
    paragraph = cell_photo.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    
    processed_photo = "temp_passport_photo.jpg"
    if photo_path and process_passport_photo(photo_path, processed_photo):
        run.add_picture(processed_photo, width=Inches(1.2))
        if os.path.exists(processed_photo):
            os.remove(processed_photo)
    else:
        run.add_text("[No Photo]")

    doc.add_paragraph() # Spacer

    # --- 2. Dynamic Sections ---
    ignore_keys = ["Full Name", "Email", "Phone", "Address", "LinkedIn", "Photo Path"]

    for section, content in personal_info.items():
        # Skip if key is in ignore list or content is empty/NA
        if section in ignore_keys or content == "NA" or not content:
            continue

        # Add Section Heading with Border
        heading = doc.add_heading(section.upper(), level=1)
        add_bottom_border(heading)
        
        # Check if this is the Education section (which is now a list of lists)
        if section == "Education" and isinstance(content, list):
            # Create Table
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid' # Gives it the black borders
            
            # Header Row
            hdr_cells = table.rows[0].cells
            headers = ["Degree/Course", "Institution", "Year", "Grade/CGPA"]
            for i, h_text in enumerate(headers):
                run = hdr_cells[i].paragraphs[0].add_run(h_text)
                run.bold = True
            
            # Data Rows
            for edu_row in content:
                row_cells = table.add_row().cells
                row_cells[0].text = edu_row[0]
                row_cells[1].text = edu_row[1]
                row_cells[2].text = edu_row[2]
                row_cells[3].text = edu_row[3]
            
            doc.add_paragraph() # Spacer after table
            
        else:
            # Standard Paragraph for other sections (Experience, Skills, etc.)
            p = doc.add_paragraph(str(content))
            p.style = 'List Bullet' if "\n" in str(content) else 'Normal'

    # Save
    output_filename = f"{personal_info.get('Full Name', 'Resume').replace(' ', '_')}_Resume.docx"
    doc.save(output_filename)
    print(f"\nSuccess! Resume saved as: {output_filename}")

def main():
    print("--- Basic Resume Builder ---")
    data = {}

    # 1. Standard Fields
    standard_fields = ["Full Name", "Email", "Phone", "Address", "LinkedIn", "Objective"]
    for field in standard_fields:
        user_input = input(f"Enter {field}: ").strip()
        data[field] = user_input if user_input else "NA"

    # 2. Education (Special Collection for Table)
    print("\n--- Education Details ---")
    education_list = []
    while True:
        degree = input("Degree (e.g., B.Tech) [Press Enter to finish education]: ").strip()
        if not degree:
            break
        school = input("Institution/University: ").strip()
        year = input("Year of Passing: ").strip()
        grade = input("Grade/Percentage: ").strip()
        education_list.append([degree, school, year, grade])
        print("---")
    
    data["Education"] = education_list if education_list else "NA"

    # 3. Text Blocks
    text_fields = ["Experience", "Skills", "Projects"]
    print("\n(Tip: For lists, you can manually type bullets or just use new lines)")
    for field in text_fields:
        user_input = input(f"Enter {field}: ").strip()
        data[field] = user_input if user_input else "NA"

    # 4. Photo
    photo_input = input("\nEnter path to your photo (e.g., profile.jpg) or 'NA': ").strip()
    photo_path = photo_input if photo_input.lower() != "na" else None

    generate_resume(data, photo_path)

if __name__ == "__main__":
    main()